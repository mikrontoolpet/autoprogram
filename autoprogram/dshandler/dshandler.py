from docx import Document
from docx.shared import Mm
import logging

# Set docxhandler logging level to INFO
logging.basicConfig(level=logging.INFO)
_logger = logging.getLogger(__name__)


class DataSheet:
    def __init__(self, path):
        self.path = path

    def __enter__(self):
        self.document = Document()
        return self

    def __exit__(self, type, value, traceback):
        self.document.save(self.path)

    def add_heading(self, heading):
        self.document.add_heading(heading)

    def add_text_arguments(self, ds_text_args):
        for text_arg in ds_text_args:
            self.document.add_paragraph(text_arg, style='List Bullet')
        _logger.info("Text added!!!")

    def add_wheelpacks_table(self, whp_df):
        """
        whp_df is the dataframe containing wheelpacks names and their positions
        """
        table = self.document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Position'
        hdr_cells[1].text = 'Wheelpack Name'
        for row in whp_df.iterrows():
            whp_name = row[1].whp_name
            whp_posn = row[0]
            row_cells = table.add_row().cells
            row_cells[0].text = str(whp_posn)
            row_cells[1].text = str(whp_name)
        _logger.info("Wheelpacks table created!!!")

    def add_pictures(self, ds_img_paths, width, heigth):
        table = self.document.add_table(rows=1, cols=2)
        row_cells = table.add_row().cells
        for i, ds_img_path in enumerate(ds_img_paths):
            paragraph = row_cells[i].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(str(ds_img_path), width=Mm(width), height=Mm(heigth))
        _logger.info("Pictures inserted!!!")